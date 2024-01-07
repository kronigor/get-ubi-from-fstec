import pandas as pd
import json
import requests
import os
import urllib3

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Pt


def change_orientation(document):
    """ Change the orientation of all sections in a Word document to landscape. """
    for section in document.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
    return document


def save_df_to_docx(df, change_orient):
    """ Takes a pandas DataFrame and saves it as a table in a Word document """
    try:
        filename = 'UBI.docx'
        print(f"[INFO]: Create a Word document and saves it to '{filename}'")
        doc = Document()
        if change_orient:
            doc = change_orientation(doc)
        doc.add_heading('Банк данных угроз безопасности информации', level=1)
        columns = ['№ УБИ', 'Описание', 'Источники угрозы', 'Объект воздействия', 'Последствия реализации угрозы']
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        table = doc.add_table(df.shape[0] + 1, len(columns), style='Table Grid')
        for i, col in enumerate(columns):
            table.cell(0, i).text = col
            table.cell(0, i).paragraphs[0].runs[0].font.bold = True
        columns_len = len(table.columns)
        table_cells = table._cells[columns_len:]
        for i in range(df.shape[0]):
            row_cells = table_cells[i * columns_len: (i + 1) * columns_len]
            for j in range(df.shape[1]):
                if j == 0:
                    val = str(df.values[i, j])
                    row_cells[j].text = f"УБИ.{'0' * (3 - len(val))}{val}"
                else:
                    row_cells[j].text = str(df.values[i, j])
        doc.save(filename)
        print('[INFO]: Done!')
    except Exception as e:
        print(f"[ERROR]: {e}")


def parse_threatlist(change_orient):
    """ Parses the 'thrlist.xlsx' file, extracts relevant data and converts it to JSON """
    try:
        filename = 'UBI.json'
        print(f"[INFO]: Parses the 'thrlist.xlsx' file and converts it to '{filename}'")
        df = pd.read_excel('thrlist.xlsx', skiprows=1)
        for i in ('Дата включения угрозы в БнД УБИ', 'Дата последнего изменения данных'):
            df[i] = df[i].dt.strftime('%d-%m-%Y')
        json_file = df.to_dict('records')
        columns_posl = ('Нарушение конфиденциальности', 'Нарушение целостности', 'Нарушение доступности')
        for j in columns_posl:
            df[j] = df[j].replace({1: j, 0: ''})
        with open(filename, 'w', encoding='utf-8') as file:
            json.dump(json_file, file)

        df['Последствия'] = df[columns_posl[0]] + '\n' + df[columns_posl[1]] + '\n' + df[columns_posl[2]]
        df = df.drop(columns=[*columns_posl, 'Описание', 'Дата включения угрозы в БнД УБИ', 'Дата последнего '
                                                                                            'изменения данных'], axis=1)
        print('[INFO]: Done!')
        save_df_to_docx(df, change_orient)
        return
    except Exception as e:
        print(f"[ERROR]: {e}")


def get_threatlist(change_orient):
    """ Get 'thrlist.xlsx' from bdu.fstec.ru if the file is not available locally """
    try:
        urllib3.disable_warnings()
        filename = 'thrlist.xlsx'
        if not os.path.isfile(filename):
            print(f"[INFO]: Download {filename} from https://bdu.fstec.ru/")
            url = 'https://bdu.fstec.ru/files/documents/thrlist.xlsx'
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) '
                              'Chrome/117.0.0.0 Safari/537.36'}
            response = requests.get(url, headers=headers, verify=False)
            with open(filename, 'wb') as file:
                file.write(response.content)
            print('[INFO]: Done!')
        else:
            print(f"[INFO]: '{filename}' exists locally.")
        parse_threatlist(change_orient)
        return
    except Exception as e:
        print(f"[ERROR]: {e}")


if __name__ == "__main__":
    try:
        get_threatlist(change_orient=True)
    except Exception as e:
        print(e)
